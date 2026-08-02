"""Global Markets Daily Brief - DOCX 报告生成器（10 分钟可读版）

设计标准（硬约束）：
- 正文 <= 8 页，TL;DR 严格 1 页，每大节最多 1 张表 + 少量点评
- 同类行情统一表格模板（右对齐 / 跌红涨绿 / 斑马纹），禁止 mono dump
- 真标题层级（Heading 1/2）+ Word TOC 域活目录
- 字号四级体系：H1 16 / H2 12 / 正文 10 / 脚注 8
- 工程状态（质量告警/解析失败/内部标记）与正文彻底隔离，统一进末页附录
- 一处数据只出现一次；超预算明细下沉附录
"""
import os
import re
from datetime import datetime, timezone, timedelta

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from storage import db


# ── 排版常量（字号 = 信息优先级，四级体系）─────────────────
FONT_TITLE = 24
FONT_H1 = 16
FONT_H2 = 12
FONT_BODY = 10
FONT_TABLE = 9
FONT_CAPTION = 8

# ── 颜色方案 ──────────────────────────────────────────────
BLUE_DARK = RGBColor(0x1B, 0x3A, 0x5C)
BLUE_MED = RGBColor(0x2E, 0x75, 0xB6)
GRAY = RGBColor(0x80, 0x80, 0x80)
GREEN = RGBColor(0x00, 0x80, 0x00)
RED = RGBColor(0xCC, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# ── 每节信息预算（10 分钟红线）────────────────────────────
BUDGET_LINES = {
    "news_flow": 6,
    "positioning": 6,
    "credit_rates": 6,
    "macro": 8,
    "central_bank": 8,
    "talking_points": 6,
    "asia_day_ahead": 8,
    "key_levels": 14,
}

# 工程噪声告警（不进正文，仅附录小字）
_NOISE_RE = re.compile(
    r"settlement fallback|close==prev_close|no change\b|unconfirmed|"
    r"tickers not resolved|not available|疑似数据源未更新",
    re.IGNORECASE,
)

# 需要从正文剥离/丢弃的内部标记
_STRIP_PREFIX_RE = re.compile(r"^\s*\[BBG-EXCLUSIVE\]\s*", re.IGNORECASE)
_DROP_LINE_RE = re.compile(
    r"unconfirmed|tickers not resolved|level not available|not resolved in current feed",
    re.IGNORECASE,
)


# ══════════════════════════════════════════════════════════
# 基础构件
# ══════════════════════════════════════════════════════════

def _set_cell_shading(cell, color_hex: str):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_text(cell, text: str, bold=False, size=FONT_TABLE, color=BLACK,
                   alignment=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


def _add_section_header(doc, title: str, level=1):
    """真标题层级：H1=16pt 深蓝，H2=12pt 中蓝（进入 Word 导航窗格）"""
    h = doc.add_heading(title, level=level)
    # 节间改为自然流动（去掉硬分页），用段前间距保持视觉区分
    h.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    h.paragraph_format.space_after = Pt(4)
    for run in h.runs:
        run.font.color.rgb = BLUE_DARK if level == 1 else BLUE_MED
        run.font.size = Pt(FONT_H1 if level == 1 else FONT_H2)
    return h


def _add_toc_field(doc):
    """插入可更新的 Word TOC 域（活目录，可跳转/右键更新）"""
    p = doc.add_paragraph()
    fld = parse_xml(
        f'<w:fldSimple {nsdecls("w")} w:instr=" TOC \\o &quot;1-2&quot; \\h \\z \\u ">'
        f'<w:r><w:rPr><w:i/><w:color w:val="808080"/><w:sz w:val="{FONT_CAPTION*2}"/></w:rPr>'
        f"<w:t>Table of Contents \u2014 right-click \u2192 Update Field to refresh</w:t></w:r>"
        f"</w:fldSimple>"
    )
    p._p.append(fld)


def _caption(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(FONT_CAPTION)
    run.font.italic = True
    run.font.color.rgb = GRAY


def _so_what(doc, text: str):
    """每节开头的一句 bold 结论句"""
    if not text:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(FONT_BODY)


def _body_par(doc, text: str, size=FONT_BODY, bold=False, color=BLACK, bullet=False):
    p = doc.add_paragraph(style="List Bullet" if bullet else None)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    return p


_NUMERIC_RE = re.compile(r"^[+\-\u2212\u2248~]?[$\u20ac\u00a3\u00a5]?[\d,.]+")


def _chg_color(text: str):
    s = str(text).strip()
    if s.startswith(("+",)):
        return GREEN
    if s.startswith(("-", "\u2212")):
        return RED
    return BLACK


def _add_market_table(doc, headers: list, rows: list, col_widths: list = None):
    """统一行情表格：表头深蓝白字、斑马纹、数字右对齐、涨绿跌红"""
    if not rows:
        _body_par(doc, "Data unavailable.", size=FONT_TABLE, color=GRAY)
        return None
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    chg_cols = {i for i, h in enumerate(headers)
                if re.search(r"chg|%|bps|change", str(h), re.IGNORECASE)}

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_shading(cell, "1B3A5C")
        _set_cell_text(cell, h, bold=True, size=FONT_TABLE, color=WHITE,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER)

    for r_idx, row_data in enumerate(rows):
        for c_idx in range(len(headers)):
            val = row_data[c_idx] if c_idx < len(row_data) else ""
            cell = table.rows[r_idx + 1].cells[c_idx]
            if r_idx % 2 == 1:
                _set_cell_shading(cell, "EDF2F9")
            text = str(val) if val is not None else ""
            numeric = bool(_NUMERIC_RE.match(text))
            color = _chg_color(text) if c_idx in chg_cols else BLACK
            _set_cell_text(
                cell, text, size=FONT_TABLE, color=color,
                alignment=WD_ALIGN_PARAGRAPH.RIGHT if numeric else WD_ALIGN_PARAGRAPH.LEFT,
            )

    if col_widths and len(col_widths) == len(headers):
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    else:
        total = 16.0
        w = total / len(headers)
        for row in table.rows:
            for i in range(len(headers)):
                row.cells[i].width = Cm(w)
    return table


# ══════════════════════════════════════════════════════════
# 内容清洗与预算渲染
# ══════════════════════════════════════════════════════════

def _clean_lines(text: str) -> list:
    """剥离内部标记、丢弃失败/空行，返回可呈现行"""
    out = []
    for ln in (text or "").splitlines():
        ln = _STRIP_PREFIX_RE.sub("", ln.rstrip())
        s = ln.strip()
        if not s:
            continue
        if _DROP_LINE_RE.search(s):
            continue
        # 空标签行（如 "Sovereign CDS 5Y:"）
        if s.endswith(":") and len(s) <= 40:
            continue
        out.append(s)
    return out


def _is_shout_header(s: str) -> bool:
    """全大写短行（如 WHERE IS THE CROWD WRONG?）视为小节标题"""
    letters = [c for c in s if c.isalpha()]
    if not letters or len(s) > 60:
        return False
    upper = sum(1 for c in letters if c.isupper())
    return upper / len(letters) > 0.8


# ── 可扫读性：粗体导语 + 短正文（Nielsen 扫描研究 / F-pattern）──
_SENT_SPLIT = re.compile(r"(?<=[.!?])\s+(?=[A-Z\u201c\"(0-9~])")

_ACRONYMS = {
    "US", "EU", "UK", "AI", "FED", "FOMC", "ECB", "BOJ", "RBA", "PBOC", "PBoC",
    "MLIV", "GMT", "EDT", "HKT", "KOSPI", "CPI", "PPI", "GDP", "EM", "FX",
    "HY", "IG", "OAS", "CFTC", "OIS", "USD", "JPY", "EUR", "CNH", "CNY",
    "USDJPY", "USDCNH", "EURUSD", "YTD", "QT", "QE", "MAS", "RBNZ", "BOE",
    "SNB", "OPEC", "IMF", "ETF", "10Y", "2Y", "30Y", "5Y", "MSCI", "CSI",
}


def _smart_title(s: str) -> str:
    """全大写标题转 Title Case（保留已知缩写），降低'喊叫感'"""
    letters = [c for c in s if c.isalpha()]
    if not letters or sum(1 for c in letters if c.isupper()) / len(letters) < 0.7:
        return s  # 已是混合大小写，不动
    out = []
    for w in s.split():
        core = re.sub(r"[^A-Za-z0-9]", "", w)
        if core.upper() in _ACRONYMS:
            out.append(w)
        elif core.isupper():
            out.append(w.capitalize())
        else:
            out.append(w)
    return " ".join(out)


def _split_lead(s: str):
    """把长条目拆为 (导语, 正文)。导语加粗做扫描锚点，正文常规。

    识别三种形态：
    1. "~04:00 GMT — HEADLINE WORDS: body..."  （News Flow 时间戳条目）
    2. "Label/Claim: body..."                   （冒号前短标签）
    3. "First short sentence. Rest..."          （Talking Points 观点句）
    """
    m = re.match(
        r"^([~\u2248]?\s*\d{1,2}:\d{2}\s*(?:GMT|EDT|HKT|BST)?[^:]*?\u2014\s*[^:]+?):\s+(.+)$",
        s)
    if m:
        lead, body = m.group(1).strip(), m.group(2).strip()
        # 双冒号标题：正文若以全大写子标题开头（"ECONOMY COOLING, ...:"）并入导语
        m2 = re.match(r"^([A-Z0-9 ,'&/%\-\.\$\u00a5\u2013\u2014?]{8,80}):\s+(.+)$", body)
        if m2 and _is_shout_header(m2.group(1)[:60]):
            lead = lead + ": " + m2.group(1).strip()
            body = m2.group(2).strip()
        return lead, body
    i = s.find(": ")
    if 0 < i <= 75 and not s[:i].rstrip().endswith(("e.g", "i.e", "vs")):
        return s[:i].strip(), s[i + 2:].strip()
    parts = _SENT_SPLIT.split(s, 1)
    if len(parts) == 2 and len(parts[0].split()) <= 22:
        return parts[0].strip(), parts[1].strip()
    return None, s


def _trim_body(text: str, max_sent=2, max_words=48):
    """正文压到 <=2 句 / ~48 词；返回 (截断文本, 是否截断)"""
    sents = _SENT_SPLIT.split(text or "")
    out, n = [], 0
    for x in sents:
        w = len(x.split())
        if out and (len(out) >= max_sent or n + w > max_words):
            break
        out.append(x)
        n += w
    trimmed = " ".join(out).strip()
    return trimmed, len(trimmed) < len((text or "").strip()) - 15


def _render_budgeted(doc, section, budget: int, overflow_sink: list,
                     section_name: str, lead_bold=True, numbered=False):
    """按信息预算渲染章节；超出部分下沉附录 overflow_sink"""
    if not section:
        return
    t = section.get("type")
    if t == "table":
        _add_market_table(doc, section["headers"], section["rows"])
        for note in _clean_lines("\n".join(section.get("notes", [])))[:3]:
            _body_par(doc, note, size=FONT_TABLE)
        return
    if t == "bullets":
        lines = _clean_lines("\n".join(section.get("items", [])))
    else:
        lines = _clean_lines(section.get("text", ""))

    shown, dropped = lines[:budget], lines[budget:]
    n = 0
    any_trimmed = False
    for s in shown:
        if _is_shout_header(s):
            _body_par(doc, s.rstrip("?").title() + ("?" if s.endswith("?") else ""),
                      size=FONT_BODY, bold=True, color=BLUE_MED)
            continue
        n += 1
        s = re.sub(r"^\d+\.\s*", "", s)  # 去掉源数据自带的手敲编号
        lead, rest = _split_lead(s)
        body, was_trimmed = _trim_body(rest) if rest else ("", False)
        if was_trimmed:
            any_trimmed = True
            dropped.append(s)  # 全文进附录备查

        p = doc.add_paragraph(style=None if numbered else "List Bullet")
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(9)
        prefix = f"{n}. " if numbered else ""
        if lead:
            r = p.add_run(prefix + _smart_title(lead)
                          + (" \u2014 " if body and not lead.endswith((".", "?", "!")) else " "))
            r.font.size = Pt(FONT_BODY)
            r.font.bold = True
            r.font.color.rgb = BLUE_DARK
        elif prefix:
            r = p.add_run(prefix)
            r.font.size = Pt(FONT_BODY)
            r.font.bold = True
        if body or not lead:
            r2 = p.add_run(body if lead else s)
            r2.font.size = Pt(FONT_BODY)
            r2.font.color.rgb = BLACK
            if was_trimmed:
                r3 = p.add_run(" \u2026")
                r3.font.size = Pt(FONT_BODY)
                r3.font.color.rgb = GRAY
    if dropped:
        note = f"(+{len(dropped)} more items"
        if any_trimmed:
            note = f"(details trimmed; +{len(dropped)} full items"
        _body_par(doc, note + " \u2014 see Appendix)", size=FONT_CAPTION, color=GRAY)
        overflow_sink.append((section_name, dropped))


def add_trend_icon(change_pct: float) -> str:
    if change_pct > 0.5:
        return "\u25b2"
    elif change_pct < -0.5:
        return "\u25bc"
    return "\u2192"


def _parse_pct(s: str):
    """解析涨跌幅（支持 % 与 bps；bps 折算为可比较的量级）"""
    txt = str(s)
    m = re.search(r"([+\-\u2212]?\d+(?:\.\d+)?)\s*%", txt)
    if m:
        return float(m.group(1).replace("\u2212", "-"))
    m = re.search(r"([+\-\u2212]?\d+(?:\.\d+)?)\s*bps", txt, re.IGNORECASE)
    if m:
        return float(m.group(1).replace("\u2212", "-")) / 10.0
    return None


def _tldr_rows_and_bullets(bbg_sections: dict, market_data: dict, fx_rates: dict):
    """构建 TL;DR 核心总表 + 自动要点（每资产类挑最大变动）"""
    rows, bullets = [], []
    label = {"equities": "Equities", "rates": "Rates", "commodities": "Commodities", "fx": "FX"}
    for key in ["equities", "rates", "commodities", "fx"]:
        sec = bbg_sections.get(key)
        if not (sec and sec.get("type") == "table"):
            continue
        best, best_abs = None, -1.0
        for r in sec["rows"][:8]:
            if len(r) < 4:
                continue
            pct = _parse_pct(r[3])
            if pct is not None and abs(pct) > best_abs:
                best_abs, best = abs(pct), (r[0], r[3])
        # 总表：每类取前 2 行
        for r in sec["rows"][:2]:
            if len(r) >= 4:
                rows.append([label[key], r[0], r[2], r[3]])
        if best:
            bullets.append(f"{label[key]}: {best[0]} {best[1]} led the move.")
    # 无 BBG 时回退免费源
    if not rows:
        for cat, lbl in [("indices", "Equities"), ("rates", "Rates"), ("commodities", "Commodities")]:
            for sym, info in list(market_data.get(cat, {}).items())[:2]:
                rows.append([lbl, info.get("name", sym), str(info.get("price", "")),
                             f"{info.get('change_pct', 0):+.2f}%"])
        for pair, info in list(fx_rates.items())[:2]:
            rows.append(["FX", pair, str(info.get("rate", "")),
                         f"{float(info.get('change_pct', 0) or 0):+.2f}%"])
    return rows, bullets[:4]


def _split_alerts(quality_alerts: list, crosscheck: dict):
    """告警分流：实质冲突 vs 工程噪声（全部只进附录）"""
    substantive, noise = [], []
    for a in (quality_alerts or []):
        (noise if _NOISE_RE.search(str(a)) else substantive).append(str(a))
    cc = [str(a) for a in (crosscheck or {}).get("alerts", [])]
    return substantive, noise, cc


# ══════════════════════════════════════════════════════════
# 主入口
# ══════════════════════════════════════════════════════════

def generate_report(
    market_data: dict,
    economic_data: dict,
    fx_rates: dict,
    news_articles: list,
    calendar_events: list,
    output_path: str,
    quality_alerts: list = None,
    bbg: dict = None,
    crosscheck: dict = None,
) -> str:
    """生成 Global Markets Daily Brief DOCX（签名与旧版一致）"""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(FONT_BODY)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    now = datetime.now(timezone.utc) + timedelta(hours=8)
    date_display = now.strftime("%A, %B %d, %Y")
    time_display = now.strftime("%I:%M %p").lstrip("0").lower()

    bbg_sections = (bbg or {})
    overflow_sink = []  # [(section_name, [lines])] 附录明细

    # ════════════════════════════════════════════════════
    # PAGE 1 — 标题块（紧凑）+ TL;DR（严格 1 页）
    # ════════════════════════════════════════════════════
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("GLOBAL MARKETS DAILY BRIEFING")
    run.font.size = Pt(FONT_TITLE)
    run.font.bold = True
    run.font.color.rgb = BLUE_DARK

    dateline = doc.add_paragraph()
    dateline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = dateline.add_run(f"{date_display} \u00b7 Hong Kong {time_display} HKT")
    run.font.size = Pt(FONT_BODY)

    src_note = doc.add_paragraph()
    src_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = src_note.add_run(
        "PRIMARY: Bloomberg ASKB \u00b7 Corroborated: Yahoo / FRED / Alpha Vantage"
        if bbg else "Sources: Yahoo Finance \u00b7 FRED \u00b7 Alpha Vantage (no Bloomberg input)"
    )
    run.font.size = Pt(FONT_CAPTION)
    run.font.italic = True
    run.font.color.rgb = GRAY

    # ── TL;DR / Executive Summary ──
    _add_section_header(doc, "Executive Summary (TL;DR)")
    tldr_rows, tldr_bullets = _tldr_rows_and_bullets(bbg_sections, market_data, fx_rates)

    # 头条（清洗后 news_flow 第一行）
    nf = bbg_sections.get("news_flow")
    if nf:
        nf_lines = _clean_lines(nf.get("text", "") if nf.get("type") == "text"
                                else "\n".join(nf.get("items", [])))
        headline = next((l for l in nf_lines if not _is_shout_header(l)), None)
        if headline:
            tldr_bullets.insert(0, f"Top story: {headline}")

    for b in tldr_bullets[:5]:
        _body_par(doc, b, size=FONT_BODY, bullet=True)

    if tldr_rows:
        _add_market_table(doc, ["Class", "Asset", "Last", "Chg"], tldr_rows[:8],
                          col_widths=[3.0, 5.0, 4.0, 4.0])
        _caption(doc, "Cross-asset snapshot \u00b7 full tables in Section 3")

    # ── 活目录（TOC 域）──
    _add_toc_field(doc)
    disclaimer = doc.add_paragraph()
    run = disclaimer.add_run(
        "Prepared for educational and interview preparation purposes only. "
        "Does not constitute investment advice."
    )
    run.font.size = Pt(FONT_CAPTION)
    run.font.italic = True
    run.font.color.rgb = GRAY

    doc.add_page_break()

    # ════════════════════════════════════════════════════
    # SECTION 1 + 2 — 日历与今日事件（同页）
    # ════════════════════════════════════════════════════
    _add_section_header(doc, "Section 1: Monthly Macro Calendar (past 7d + upcoming)")
    today = now.strftime("%Y-%m-%d")

    def _cal_status(ev):
        d = ev["date"][:10]
        if d == today:
            return "Today"
        if d < today:
            return "Past"
        return "Upcoming"

    # 过去一周 + 今日永远展示；其余预算补未来事件（避免刚发生的会议被截断）
    past_today = [ev for ev in calendar_events if _cal_status(ev) in ("Past", "Today")]
    upcoming = [ev for ev in calendar_events if _cal_status(ev) == "Upcoming"]
    CAL_BUDGET = 16
    keep_upcoming = max(0, CAL_BUDGET - len(past_today))
    shown_events = past_today + upcoming[:keep_upcoming]
    overflow_events = upcoming[keep_upcoming:]
    cal_rows = [[ev["date"], ev["event"], ev["country"], _cal_status(ev)]
                for ev in shown_events]
    if overflow_events:
        overflow_sink.append(("Macro Calendar",
                              [" | ".join([ev["date"], ev["event"], ev["country"]])
                               for ev in overflow_events]))
    _add_market_table(doc, ["Date", "Event", "Country", "Status"], cal_rows,
                      col_widths=[2.5, 8, 3, 2.5])
    if overflow_events:
        _caption(doc,
                 f"Showing past 7d + next {len(upcoming[:keep_upcoming])} events "
                 f"\u00b7 rest in Appendix")

    _add_section_header(doc, "Section 2: Today's Events", level=2)
    today_events = [e for e in calendar_events if e["date"][:10] == now.strftime("%Y-%m-%d")]
    if today_events:
        today_rows = [[ev["event"], ev["country"], "\u2605" * ev.get("importance", 3)]
                      for ev in today_events]
        _add_market_table(doc, ["Event", "Country", "Significance"], today_rows,
                          col_widths=[8, 3, 3])
    else:
        _body_par(doc, "No major events scheduled for today.")

    # ════════════════════════════════════════════════════
    # SECTION 3 — 隔夜复盘（统一表格，无告警墙、无 mono）
    # ════════════════════════════════════════════════════
    _add_section_header(doc, "Section 3: Overnight Market Recap")

    # So-what 结论句：由各类最大变动合成
    movers = [b.split(": ", 1)[1].replace(" led the move.", "")
              for b in tldr_bullets
              if b.startswith(("Equities", "Rates", "Commodities", "FX")) and ": " in b]
    if movers:
        _so_what(doc, "So what: biggest overnight moves \u2014 " + ", ".join(movers[:4]) + ".")

    _sub_specs = [
        ("equities", "Equities", "indices"),
        ("rates", "Fixed Income", "rates"),
        ("commodities", "Commodities", "commodities"),
        ("fx", "Foreign Exchange", None),
    ]
    for key, sub_title, md_key in _sub_specs:
        _add_section_header(doc, sub_title, level=2)
        sec = bbg_sections.get(key)
        if sec and sec.get("type") == "table":
            _add_market_table(doc, sec["headers"], sec["rows"])
            for note in _clean_lines("\n".join(sec.get("notes", [])))[:2]:
                _body_par(doc, note, size=FONT_TABLE)
            _caption(doc, "Source: Bloomberg ASKB (primary)")
        elif sec:
            # 解析失败兜底：清洗后按正文行呈现（不再 mono dump）
            _render_budgeted(doc, sec, 8, overflow_sink, sub_title, lead_bold=False)
            _caption(doc, "Source: Bloomberg ASKB (primary)")
        elif key == "fx" and fx_rates:
            fx_rows = []
            for pair, info in fx_rates.items():
                try:
                    chg = float(info.get("change_pct", 0) or 0)
                except (ValueError, TypeError):
                    chg = 0
                fx_rows.append([pair, str(info.get("rate", "")), f"{chg:+.2f}%",
                                add_trend_icon(chg)])
            _add_market_table(doc, ["Pair", "Rate", "Chg %", "Trend"], fx_rows,
                              col_widths=[3, 3, 2.5, 1.5])
        elif md_key:
            rows = []
            for sym, info in market_data.get(md_key, {}).items():
                rows.append([info.get("name", sym), str(info.get("price", "")),
                             f"{info.get('change_pct', 0):+.2f}%", info.get("trend", "\u2192")])
            _add_market_table(doc, ["Asset", "Last", "Chg %", "Trend"], rows,
                              col_widths=[5, 3, 2.5, 1.5])

    # Positioning / Credit（BBG 独家，预算渲染）
    if bbg_sections.get("positioning"):
        _add_section_header(doc, "Positioning & Flows", level=2)
        _render_budgeted(doc, bbg_sections["positioning"], BUDGET_LINES["positioning"],
                         overflow_sink, "Positioning & Flows")
        _caption(doc, "Source: Bloomberg ASKB (primary \u00b7 not externally verifiable)")
    if bbg_sections.get("credit_rates"):
        _add_section_header(doc, "Credit & Rates Microstructure", level=2)
        _render_budgeted(doc, bbg_sections["credit_rates"], BUDGET_LINES["credit_rates"],
                         overflow_sink, "Credit & Rates Microstructure")
        _caption(doc, "Source: Bloomberg ASKB (primary \u00b7 not externally verifiable)")

    # ════════════════════════════════════════════════════
    # SECTION 4 + 5 — 宏观动态 / 央行观察（同页）
    # ════════════════════════════════════════════════════
    _add_section_header(doc, "Section 4: Latest Macro Developments")
    if bbg_sections.get("news_flow"):
        _add_section_header(doc, "Markets Live \u2014 News Flow", level=2)
        _render_budgeted(doc, bbg_sections["news_flow"], BUDGET_LINES["news_flow"],
                         overflow_sink, "News Flow")
        _caption(doc, "Source: Bloomberg ASKB \u2014 FIRST WORD / TOP / MLIV (primary)")

    if bbg_sections.get("macro"):
        _render_budgeted(doc, bbg_sections["macro"], BUDGET_LINES["macro"],
                         overflow_sink, "Macro Developments")
        _caption(doc, "Source: Bloomberg ASKB (primary)")
    elif economic_data:
        eco_rows = []
        for sid, info in economic_data.items():
            if not sid.endswith("_prev"):
                eco_rows.append([info.get("label", sid), info.get("value", "N/A"),
                                 info.get("prev_value", "N/A"),
                                 info.get("reference_period", info.get("date", ""))])
        _add_market_table(doc, ["Indicator", "Latest", "Previous", "Period"],
                          eco_rows[:10], col_widths=[5.5, 2.5, 2.5, 2.5])
    else:
        _body_par(doc, "Economic data feed currently unavailable.", color=GRAY)

    _add_section_header(doc, "Section 5: Central Bank Monitor")
    if bbg_sections.get("central_bank"):
        _render_budgeted(doc, bbg_sections["central_bank"], BUDGET_LINES["central_bank"],
                         overflow_sink, "Central Bank Monitor")
        _caption(doc, "Source: Bloomberg ASKB (primary)")
    else:
        cb_events = [e for e in calendar_events if e.get("type") == "central_bank"]
        if cb_events:
            cb_rows = []
            for ev in cb_events:
                expected = db.get_expected_central_bank_rate(ev["country"], ev["date"])
                cb_rows.append([ev["country"], ev["event"], ev["date"], expected])
            _add_market_table(doc, ["Central Bank", "Event", "Date", "Expected"],
                              cb_rows, col_widths=[3, 6.5, 2.5, 2.5])
        else:
            _body_par(doc, "No central bank events in the near term.")

    # ════════════════════════════════════════════════════
    # SECTION 6 — 访谈要点
    # ════════════════════════════════════════════════════
    _add_section_header(doc, "Section 6: Interview Talking Points")
    tp_sec = bbg_sections.get("talking_points")
    if tp_sec:
        _render_budgeted(doc, tp_sec, BUDGET_LINES["talking_points"], overflow_sink,
                         "Talking Points", lead_bold=False, numbered=True)
        _caption(doc, "Source: Bloomberg ASKB (primary)")
    else:
        _ecb_exp = db.get_expected_central_bank_rate("Eurozone", "2026-07-23")
        _ecb_rate = _ecb_exp.split()[0] if _ecb_exp and _ecb_exp != "TBD" else "current levels"
        talking_points = [
            "FOMC Decision (Jul 28-29): Market pricing 85% probability of a 25bp hold. "
            "Focus on Powell's language around September.",
            f"ECB Meeting (Jul 23): Expected to hold at {_ecb_rate} (per Central Bank Monitor).",
            "BOJ Outlook Report (Jul 31): Whether BOJ signals another hike; "
            "USD/JPY 162+ keeps MOF intervention risk live.",
            "US PCE (Jul 25): Core PCE forecast +0.2% MoM; at/below consensus cements "
            "September cut expectations.",
            "Geopolitical: Hormuz tensions keep Brent elevated; primary tail risk for crude.",
        ]
        for i, tp in enumerate(talking_points, 1):
            _body_par(doc, f"{i}. {tp}")

    # ════════════════════════════════════════════════════
    # SECTION 7 — 亚洲前瞻 + Key Levels（同页起）
    # ════════════════════════════════════════════════════
    _add_section_header(doc, "Section 7: Day Ahead \u2014 Asia Session Focus")
    if bbg_sections.get("asia_day_ahead"):
        _render_budgeted(doc, bbg_sections["asia_day_ahead"], BUDGET_LINES["asia_day_ahead"],
                         overflow_sink, "Asia Day Ahead")
        _caption(doc, "Source: Bloomberg ASKB (primary)")
    else:
        asia_rows = []
        asia_data = market_data.get("indices", {})
        for sym, name, view in [
            ("000300.SS", "China (CSI 300)", "Stimulus hopes; property in focus"),
            ("^N225", "Japan (Nikkei 225)", "USD/JPY dynamics in focus"),
            ("^HSI", "Hong Kong (HSI)", "Tech and property momentum"),
        ]:
            if sym in asia_data:
                asia_rows.append([name, view, f"Current: {asia_data[sym].get('price', 'N/A')}"])
        if asia_rows:
            _add_market_table(doc, ["Market", "Session Outlook", "Level"], asia_rows,
                              col_widths=[4, 8, 4])
        else:
            _body_par(doc, "Asia market data currently unavailable.", color=GRAY)

    # ── Key Levels：只保留 BBG 支撑/阻力等新增信息（不重复 Section 3 行情）──
    kl = bbg_sections.get("key_levels")
    if kl:
        _add_section_header(doc, "Quick Reference: Key Levels (Bloomberg)")
        _render_budgeted(doc, kl, BUDGET_LINES["key_levels"], overflow_sink,
                         "Key Levels", lead_bold=False)
        _caption(doc, "Source: Bloomberg ASKB (primary) \u00b7 spot prices in Section 3, not repeated here")

    # ════════════════════════════════════════════════════
    # APPENDIX — 工程日志与超预算明细（与正文隔离）
    # ════════════════════════════════════════════════════
    substantive, noise, cc_alerts = _split_alerts(quality_alerts, crosscheck)
    if substantive or noise or cc_alerts or overflow_sink:
        doc.add_page_break()
        _add_section_header(doc, "Appendix: Data Quality & Reconciliation Log")
        _body_par(doc, "Engineering log \u2014 for data auditing only; "
                       "does not affect the narrative above.",
                  size=FONT_CAPTION, color=GRAY)

        if substantive:
            _add_section_header(doc, "Material discrepancies", level=2)
            for a in substantive:
                _body_par(doc, a, size=FONT_CAPTION, color=RED, bullet=True)
        if cc_alerts:
            _add_section_header(doc, "Cross-source reconciliation flags", level=2)
            for a in cc_alerts:
                _body_par(doc, a, size=FONT_CAPTION, color=BLUE_DARK, bullet=True)
        if noise:
            _add_section_header(doc, "Pipeline notices (informational)", level=2)
            for a in noise:
                _body_par(doc, a, size=FONT_CAPTION, color=GRAY, bullet=True)
        if overflow_sink:
            _add_section_header(doc, "Extended detail (over budget)", level=2)
            for name, lines in overflow_sink:
                _body_par(doc, name, size=FONT_CAPTION, bold=True, color=GRAY)
                for ln in lines:
                    _body_par(doc, ln, size=FONT_CAPTION, color=GRAY, bullet=True)

    # ── 保存 ──
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"  [Report] Saved to {output_path}")
    return output_path
