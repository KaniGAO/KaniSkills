"""Global Markets Daily Brief - DOCX 报告生成器

匹配样本格式：月度日历 → 今日事件 → 隔夜复盘 → 宏观动态 → 央行观察 → 访谈要点 → 亚洲前瞻 → 数据仪表盘
"""
import os
from datetime import datetime, timezone, timedelta

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from storage import db


# ── 颜色方案 ──────────────────────────────────────────────
BLUE_DARK = RGBColor(0x1B, 0x3A, 0x5C)
BLUE_MED = RGBColor(0x2E, 0x75, 0xB6)
BLUE_LIGHT = RGBColor(0xD5, 0xE8, 0xF0)
GRAY = RGBColor(0x80, 0x80, 0x80)
GREEN = RGBColor(0x00, 0x80, 0x00)
RED = RGBColor(0xCC, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def _set_cell_shading(cell, color_hex: str):
    """设置单元格背景色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_text(cell, text: str, bold=False, size=9, color=BLACK, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """设置单元格文字格式"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = "Arial"
    # 设置中文字体
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


def _add_section_header(doc, title: str, level=1):
    """添加章节标题"""
    if level == 1:
        h = doc.add_heading(title, level=1)
        for run in h.runs:
            run.font.color.rgb = BLUE_DARK
            run.font.size = Pt(16)
    elif level == 2:
        h = doc.add_heading(title, level=2)
        for run in h.runs:
            run.font.color.rgb = BLUE_MED
            run.font.size = Pt(13)


def _add_table(doc, headers: list, rows: list[list], col_widths: list[int] = None):
    """添加格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_shading(cell, "1B3A5C")
        _set_cell_text(cell, h, bold=True, size=8, color=WHITE, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # 数据行
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            if r_idx % 2 == 1:
                _set_cell_shading(cell, "EDF2F9")
            text = str(val) if val is not None else ""
            _set_cell_text(cell, text, size=8)

    # 列宽
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    return table


def add_trend_icon(change_pct: float) -> str:
    """趋势图标"""
    if change_pct > 0.5:
        return "▲"
    elif change_pct < -0.5:
        return "▼"
    return "→"


def generate_report(
    market_data: dict,
    economic_data: dict,
    fx_rates: dict,
    news_articles: list,
    calendar_events: list,
    output_path: str,
    quality_alerts: list = None,
) -> str:
    """生成完整的 Global Markets Daily Brief DOCX 报告"""
    doc = Document()

    # ── 文档样式 ──
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    now = datetime.now(timezone.utc) + timedelta(hours=8)
    date_display = now.strftime("%A, %B %d, %Y")
    time_display = now.strftime("%I:%M %p").lstrip("0").lower()

    # ════════════════════════════════════════════════════════════
    # TITLE PAGE HEADER
    # ════════════════════════════════════════════════════════════
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("GLOBAL MARKETS")
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = BLUE_DARK

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("DAILY BRIEFING")
    run.font.size = Pt(18)
    run.font.color.rgb = BLUE_MED

    edition = doc.add_paragraph()
    edition.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = edition.add_run("S&T Interview Prep Edition")
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = GRAY

    dateline = doc.add_paragraph()
    dateline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = dateline.add_run(f"{date_display}\nHong Kong {time_display} HKT")
    run.font.size = Pt(10)
    run.font.color.rgb = BLACK

    doc.add_paragraph()  # spacing

    # ── 免责声明 ──
    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disclaimer.add_run(
        "Prepared for educational and interview preparation purposes only. "
        "Does not constitute investment advice."
    )
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = GRAY

    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ════════════════════════════════════════════════════════════
    _add_section_header(doc, "Table of Contents")
    toc_items = [
        "Section 1: Monthly Macro Calendar",
        "Section 2: Today's Events",
        "Section 3: Overnight Market Recap",
        "Section 4: Latest Macro Developments",
        "Section 5: Central Bank Monitor",
        "Section 6: Interview Talking Points",
        "Section 7: Day Ahead — Asia Session Focus",
        "Quick Reference: Key Levels Dashboard",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item, style="List Number")
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.color.rgb = BLUE_MED

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # SECTION 1: MONTHLY MACRO CALENDAR
    # ════════════════════════════════════════════════════════════
    _add_section_header(doc, "Section 1: Monthly Macro Calendar")
    p = doc.add_paragraph("Event tracking list for July—September 2026.")
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.italic = True

    cal_headers = ["Date", "Event", "Country", "Status"]
    cal_rows = []
    for ev in calendar_events:
        status = "🔔 Today" if ev["date"][:10] == now.strftime("%Y-%m-%d") else "⏳ Upcoming"
        cal_rows.append([ev["date"], ev["event"], ev["country"], status])

    _add_table(doc, cal_headers, cal_rows, col_widths=[2.5, 8, 3, 3])

    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════
    # SECTION 2: TODAY'S EVENTS
    # ════════════════════════════════════════════════════════════
    _add_section_header(doc, "Section 2: Today's Events")
    p = doc.add_paragraph(f"Key events for {date_display}.")
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.italic = True

    today_events = [e for e in calendar_events if e["date"][:10] == now.strftime("%Y-%m-%d")]
    if today_events:
        today_headers = ["Event", "Country", "Significance"]
        today_rows = []
        for ev in today_events:
            stars = "★" * ev.get("importance", 3)
            today_rows.append([ev["event"], ev["country"], stars])
        _add_table(doc, today_headers, today_rows, col_widths=[8, 3, 3])
    else:
        doc.add_paragraph("No major events scheduled for today.")

    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════
    # SECTION 3: OVERNIGHT MARKET RECAP
    # ════════════════════════════════════════════════════════════
    _add_section_header(doc, "Section 3: Overnight Market Recap")

    # ── 数据质量告警（来自 reconcile/quality 校验）──
    if quality_alerts:
        alert_p = doc.add_paragraph()
        run = alert_p.add_run("⚠ Data Quality Alerts")
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RED
        for a in quality_alerts:
            ap = doc.add_paragraph(f"• {a}", style="List Bullet")
            ap.runs[0].font.size = Pt(8)
            ap.runs[0].font.color.rgb = RED

    # Equities
    p = doc.add_paragraph()
    run = p.add_run("Equities")
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLUE_MED

    equity_headers = ["Index", "Last", "Chg %", "1W Chg", "Trend"]
    equity_rows = []
    for sym, info in market_data.get("indices", {}).items():
        equity_rows.append([
            info.get("name", sym),
            str(info.get("price", "")),
            f"{info.get('change_pct', 0):+.2f}%",
            f"{info.get('weekly_change_pct', 0):+.2f}%",
            info.get("trend", "→"),
        ])
    if equity_rows:
        _add_table(doc, equity_headers, equity_rows, col_widths=[4, 2.5, 2, 2.5, 1.5])
    else:
        doc.add_paragraph("Equity data unavailable.", style="Normal")

    # Rates
    p = doc.add_paragraph()
    run = p.add_run("Fixed Income")
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLUE_MED

    rate_headers = ["Security", "Yield", "1W Change", "Trend"]
    rate_rows = []
    for sym, info in market_data.get("rates", {}).items():
        rate_rows.append([
            info.get("name", sym),
            f"{info.get('price', '')}%",
            f"{info.get('weekly_change_pct', 0):+.2f}%",
            info.get("trend", "→"),
        ])
    if rate_rows:
        _add_table(doc, rate_headers, rate_rows, col_widths=[5, 2.5, 2.5, 1.5])
    else:
        doc.add_paragraph("Rates data unavailable.", style="Normal")

    # Commodities
    p = doc.add_paragraph()
    run = p.add_run("Commodities")
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLUE_MED

    comm_headers = ["Commodity", "Last", "Chg %", "Source", "Trend"]
    comm_rows = []
    for sym, info in market_data.get("commodities", {}).items():
        comm_rows.append([
            info.get("name", sym),
            f"${info.get('price', '')}",
            f"{info.get('change_pct', 0):+.2f}%",
            info.get("source", "Yahoo Finance"),
            info.get("trend", "→"),
        ])
    if comm_rows:
        _add_table(doc, comm_headers, comm_rows, col_widths=[3.5, 2.2, 2.0, 2.0, 1.3])
    else:
        doc.add_paragraph("Commodities data unavailable.", style="Normal")

    # FX
    p = doc.add_paragraph()
    run = p.add_run("Foreign Exchange")
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLUE_MED

    fx_headers = ["Pair", "Rate", "Change %", "Trend"]
    fx_rows = []
    for pair, info in fx_rates.items():
        try:
            chg = float(info.get("change_pct", 0) or 0)
        except (ValueError, TypeError):
            chg = 0
        fx_rows.append([pair, str(info.get("rate", "")), f"{chg:+.2f}%", add_trend_icon(chg)])
    # FX 主源为空时退化为 DXY（Yahoo market_data["fx"]），避免整段空白
    if not fx_rows:
        for sym, info in market_data.get("fx", {}).items():
            fx_rows.append([info.get("name", sym), str(info.get("price", "")), "—", info.get("trend", "→")])
    if fx_rows:
        _add_table(doc, fx_headers, fx_rows, col_widths=[3, 3, 2.5, 1.5])
    else:
        doc.add_paragraph("FX data unavailable.", style="Normal")

    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════
    # SECTION 4: LATEST MACRO DEVELOPMENTS
    # ════════════════════════════════════════════════════════════
    _add_section_header(doc, "Section 4: Latest Macro Developments")
    p = doc.add_paragraph("Recent economic data releases and key indicators.")
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.italic = True

    if economic_data:
        eco_headers = ["Indicator", "Latest", "Previous", "Reference Period"]
        eco_rows = []
        for sid, info in economic_data.items():
            if not sid.endswith("_prev"):
                eco_rows.append([
                    info.get("label", sid),
                    info.get("value", "N/A"),
                    info.get("prev_value", "N/A"),
                    info.get("reference_period", info.get("date", "")),
                ])
        if eco_rows:
            _add_table(doc, eco_headers, eco_rows[:12], col_widths=[5, 2.5, 2.5, 2.5])
    else:
        doc.add_paragraph("Economic data feed currently unavailable.", style="Normal")

    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════
    # SECTION 5: CENTRAL BANK MONITOR
    # ════════════════════════════════════════════════════════════
    _add_section_header(doc, "Section 5: Central Bank Monitor")
    cb_events = [e for e in calendar_events if e.get("type") == "central_bank"]
    if cb_events:
        cb_headers = ["Central Bank", "Event", "Date", "Expected"]
        cb_rows = []
        for ev in cb_events:
            expected = db.get_expected_central_bank_rate(ev["country"], ev["date"])
            cb_rows.append([ev["country"], ev["event"], ev["date"], expected])
        _add_table(doc, cb_headers, cb_rows, col_widths=[3, 6, 2.5, 2.5])
    else:
        doc.add_paragraph("No central bank events in the near term.")

    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════
    # SECTION 6: INTERVIEW TALKING POINTS
    # ════════════════════════════════════════════════════════════
    _add_section_header(doc, "Section 6: Interview Talking Points")
    p = doc.add_paragraph("Key narratives and discussion points for S&T interviews.")
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.italic = True

    # ECB 利率动态取自央行监控表（Section 5），避免与正文硬编码口径矛盾
    _ecb_exp = db.get_expected_central_bank_rate("Eurozone", "2026-07-23")
    _ecb_rate = _ecb_exp.split()[0] if _ecb_exp and _ecb_exp != "TBD" else "current levels"

    talking_points = [
        "FOMC Decision (Jul 28-29): Market pricing 85% probability of a 25bp hold. Focus on Powell's language around September — any hint of a pivot would be dollar-negative and EM-equity-positive.",
        f"ECB Meeting (Jul 23): Expected to hold rates at {_ecb_rate} (per Central Bank Monitor). Watch Lagarde's tone on the Eurozone growth-inflation tradeoff amid the Germany fiscal package uncertainty.",
        "BOJ/BOJ Outlook Report (Jul 31): The key question is whether the BOJ signals another rate hike. USD/JPY at 162+ is a concern — MOF intervention risk remains live.",
        "US PCE (Jul 25): Core PCE forecast at +0.2% MoM. A print at or below consensus would cement September cut expectations and likely steepen the 2s10s curve.",
        "Geopolitical Risk: Hormuz Strait tensions keep Brent elevated above $89. Iran-Israel proxy dynamics are the primary tail risk for crude into August.",
    ]
    for i, tp in enumerate(talking_points, 1):
        p = doc.add_paragraph(f"{i}. {tp}")
        p.runs[0].font.size = Pt(9)

    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════
    # SECTION 7: DAY AHEAD — ASIA SESSION FOCUS
    # ════════════════════════════════════════════════════════════
    _add_section_header(doc, "Section 7: Day Ahead — Asia Session Focus")
    p = doc.add_paragraph(f"Asia session outlook for {date_display}.")
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.italic = True

    asia_headers = ["Market", "Session Outlook", "Key Levels"]
    asia_rows = []
    asia_data = market_data.get("indices", {})
    if "000300.SS" in asia_data:
        csi = asia_data["000300.SS"]
        asia_rows.append([
            "China (CSI 300)",
            "Bullish bias on stimulus hopes post-PBOC LPR hold; property sector in focus",
            f"Current: {csi.get('price', 'N/A')}",
        ])
    if "^N225" in asia_data:
        nk = asia_data["^N225"]
        asia_rows.append([
            "Japan (Nikkei 225)",
            "Catch-up move expected post-holiday; watching USD/JPY dynamics",
            f"Current: {nk.get('price', 'N/A')}",
        ])
    if "^HSI" in asia_data:
        hsi = asia_data["^HSI"]
        asia_rows.append([
            "Hong Kong (HSI)",
            "Tech and property in focus; momentum from prior session",
            f"Current: {hsi.get('price', 'N/A')}",
        ])

    if asia_rows:
        _add_table(doc, asia_headers, asia_rows, col_widths=[3, 7, 3.5])
    else:
        doc.add_paragraph("Asia market data currently unavailable.")

    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════
    # KEY LEVELS DASHBOARD
    # ════════════════════════════════════════════════════════════
    _add_section_header(doc, "Quick Reference: Key Levels Dashboard")
    p = doc.add_paragraph(f"All levels as of {date_display}.")
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.italic = True
    p.runs[0].font.color.rgb = GRAY

    dash_headers = ["Asset", "Current", "1W Ago", "Trend"]
    dash_rows = []
    for category in ["indices", "rates", "commodities"]:
        for sym, info in market_data.get(category, {}).items():
            dash_rows.append([
                info.get("name", sym),
                str(info.get("price", "")),
                str(info.get("one_week_ago_price", info.get("price", ""))),
                info.get("trend", "→"),
            ])
    # FX
    for pair, info in fx_rates.items():
        try:
            chg = float(info.get("change_pct", 0) or 0)
        except (ValueError, TypeError):
            chg = 0
        dash_rows.append([pair, str(info.get("rate", "")), "—", add_trend_icon(chg)])

    if dash_rows:
        _add_table(doc, dash_headers, dash_rows[:12], col_widths=[4, 3, 3, 1.5])

    doc.add_paragraph()

    # ── Footer legend ──
    legend = doc.add_paragraph()
    run = legend.add_run("▲ = Uptrend  │  ▼ = Downtrend  │  → = Range-bound")
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = GRAY

    # ── 保存 ──
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"  [Report] Saved to {output_path}")
    return output_path
