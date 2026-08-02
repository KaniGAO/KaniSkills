"""Skill 自检：报告解析器。

从已生成的 .docx 报告里抽取「被审稿件」：
  - 数值断言（隔夜复盘各表格 + Key Levels 仪表盘；若报告降级为 "data unavailable" 则记录缺失）
  - 数据可用性（各复盘板块 / Key Levels 是否真的有数据）
  - 叙事正文（Section 4-7，供 agent 按 rubric 做定性评判）

只读、不渲染、不调用任何 LLM API。

兼容两种报告形态：
  A. 当前 report.py 生成、含完整表格的版本；
  B. 数据采集失败、板块降级为 "X data unavailable." 的旧版本。
"""
import os
import re
from docx import Document


# 各表格的表头签名 → 段落类别
_TABLE_SIGS = {
    "equities": ["Index", "Last", "Chg %", "1W Chg", "Trend"],
    "rates": ["Security", "Yield", "1W Change", "Trend"],
    "commodities": ["Commodity", "Last", "Chg %", "Source", "Trend"],
    "fx": ["Pair", "Rate", "Change %", "Trend"],
    "key_levels": ["Asset", "Current", "1W Ago", "Trend"],
}

# 需要抽取叙事正文的章节标题关键字
_NARRATIVE_KEYWORDS = (
    "Section 4", "Section 5", "Section 6", "Section 7",
    "Latest Macro", "Central Bank", "Interview Talking", "Day Ahead",
)

# 复盘子板块在正文中的标签 → recap 类别
_RECAP_LABELS = [
    ("Equities", "equities"),
    ("Fixed Income", "rates"),
    ("Commodities", "commodities"),
    ("Foreign Exchange", "fx"),
]


def _num(s):
    """把 '5,432.10' / '+1.23%' / '$1,800.50' / '▲' 之类的字符串解析为 float。"""
    if s is None:
        return None
    t = str(s).replace(",", "").replace("$", "").replace("%", "").strip()
    t = t.replace("▲", "").replace("▼", "").replace("→", "").strip()
    if t in ("", "—", "-", "N/A", "None", "nan"):
        return None
    try:
        return float(t)
    except ValueError:
        return None


def _report_date_from_path(path: str):
    m = re.search(r"(\d{4}-\d{2}-\d{2})", os.path.basename(path))
    return m.group(1) if m else None


def _extract_dateline(doc) -> str | None:
    for p in doc.paragraphs:
        m = re.search(r"([A-Z][a-z]+day, [A-Z][a-z]+ \d{1,2}, \d{4})", p.text)
        if m:
            return m.group(1)
    return None


def _extract_availability(doc) -> dict:
    """判断各复盘板块与 Key Levels 是否真的有数据。

    逻辑：出现子板块标签后，紧随的段落若含 'unavailable' 视为缺失。
    同时统计 doc.tables 是否被本解析器识别为复盘/Key Levels 表（有表即视为有数据）。
    """
    paragraphs = [p.text for p in doc.paragraphs]
    full_text = "\n".join(paragraphs)

    # 已被表格解析覆盖的类别（有结构化数据）
    # （此处只判定「降级文案」缺失；表格存在与否由 parse_report 的 recap 长度体现）
    availability = {k: "unknown" for _, k in _RECAP_LABELS}
    availability["key_levels"] = "unknown"

    for label, key in _RECAP_LABELS:
        if label in full_text:
            # 取该标签之后的文本片段，看是否紧跟 unavailable
            after = full_text.split(label, 1)[1]
            # 下一段落的 200 字符内
            chunk = after.strip().split("\n", 1)[0][:200]
            availability[key] = "unavailable" if "unavailable" in chunk.lower() else "available"

    if "Key Levels" in full_text:
        after = full_text.split("Key Levels", 1)[1]
        chunk = after.strip().split("\n", 1)[0][:200]
        # Key Levels 有实际数据通常会有数字；只有图例则视为缺失
        has_numbers = bool(re.search(r"\d", after[:400]))
        availability["key_levels"] = "available" if has_numbers else "unavailable"

    return availability


def parse_report(path: str) -> dict:
    doc = Document(path)
    report_date = _report_date_from_path(path)

    recap = {"equities": [], "rates": [], "commodities": [], "fx": []}
    key_levels = []

    for table in doc.tables:
        rows = [[c.text.strip() for c in r.cells] for r in table.rows]
        if not rows:
            continue
        header = rows[0]
        for sec, sig in _TABLE_SIGS.items():
            if header[:len(sig)] == sig:
                _ingest(sec, rows[1:], recap, key_levels)
                break

    narrative = _extract_narrative(doc)
    availability = _extract_availability(doc)

    # 若 recap / key_levels 实际无行但 availability 标 available，纠正为 missing（空表也视为无数据）
    for k in ("equities", "rates", "commodities", "fx"):
        if not recap[k] and availability.get(k) == "available":
            availability[k] = "missing"
    if not key_levels and availability.get("key_levels") == "available":
        availability["key_levels"] = "missing"

    return {
        "report_path": path,
        "report_date": report_date,
        "dateline": _extract_dateline(doc),
        "data_availability": availability,
        "recap": recap,
        "key_levels": key_levels,
        "narrative": narrative,
    }


def _ingest(sec, rows, recap, key_levels):
    if sec == "equities":
        for r in rows:
            recap["equities"].append({
                "name": r[0], "last": _num(r[1]), "chg_pct": _num(r[2]),
                "w1_chg": _num(r[3]), "trend": r[4] if len(r) > 4 else "",
            })
    elif sec == "rates":
        for r in rows:
            recap["rates"].append({
                "name": r[0], "yield": _num(r[1]), "w1_change": _num(r[2]),
                "trend": r[3] if len(r) > 3 else "",
            })
    elif sec == "commodities":
        for r in rows:
            recap["commodities"].append({
                "name": r[0], "last": _num(r[1]), "chg_pct": _num(r[2]),
                "source": r[3] if len(r) > 3 else "", "trend": r[4] if len(r) > 4 else "",
            })
    elif sec == "fx":
        for r in rows:
            recap["fx"].append({
                "pair": r[0], "rate": _num(r[1]), "change_pct": _num(r[2]),
                "trend": r[3] if len(r) > 3 else "",
            })
    elif sec == "key_levels":
        for r in rows:
            key_levels.append({
                "asset": r[0], "current": _num(r[1]), "w1_ago": _num(r[2]),
                "trend": r[3] if len(r) > 3 else "",
            })


def _extract_narrative(doc) -> dict:
    sections = {}
    current = None
    buf = []
    for p in doc.paragraphs:
        style = (p.style.name or "") if p.style else ""
        text = p.text.strip()
        if style.startswith("Heading") and text:
            if current:
                sections[current] = "\n".join(buf).strip()
            current = text
            buf = []
        else:
            if current and text:
                buf.append(text)
    if current:
        sections[current] = "\n".join(buf).strip()
    return {k: v for k, v in sections.items() if any(kw in k for kw in _NARRATIVE_KEYWORDS)}
